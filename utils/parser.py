"""
题库解析器 - 解析 docx 格式的考题文件
支持题型：单选题、多选题、判断题、不定项选择题、案例题
"""

import re
import hashlib
import uuid
import docx
from pathlib import Path


def _sanitize_answer(text):
    """清洗答案中的不可见/干扰字符（零宽空格、BOM等），防止 Word 文档导入时的格式污染"""
    if not text:
        return ""
    text = text.replace("\u200b", "")  # zero-width space
    text = text.replace("\u200c", "")  # zero-width non-joiner
    text = text.replace("\u200d", "")  # zero-width joiner
    text = text.replace("\u200e", "")  # left-to-right mark
    text = text.replace("\u200f", "")  # right-to-left mark
    text = text.replace("\ufeff", "")  # BOM / zero-width no-break space
    text = text.replace("\ufe0f", "")  # variation selector
    text = text.replace("\u00a0", " ")  # non-breaking space → normal space
    return text

# 知识板块映射表（细化到具体模块）
CATEGORY_MAP = {
    "咨询实务1": "心理咨询会谈技术",
    "咨询实务2": "情绪调节与压力管理",
    "咨询实务3": "心理危机识别",
    "咨询实务4": "家庭教育与心理健康科普",
    "咨询实务5": "心理咨询专业伦理与相关法律规范",
    "基础理论1": "心理学导论",
    "基础理论2": "社会心理学",
    "基础理论3": "人格心理学",
    "基础理论4": "发展心理学",
    "基础理论5": "异常心理学",
    "基础理论6": "咨询心理学",
}


def _detect_case_ranges(paragraphs):
    """
    预扫描段落，检测案例分析题题号范围。
    返回: set[int] 案例题题号集合，以及 list[dict] 案例背景信息列表
          每个背景信息: {"start_num": int, "background_text": str, "range_text": str}
    """
    case_nums = set()
    case_backgrounds = []
    for text in paragraphs:
        if not text:
            continue
        if '案例分析题' in text and '不定项选择题' in text:
            range_match = re.search(r'(\d+)～(\d+)\s*题', text)
            if range_match:
                start = int(range_match.group(1))
                end = int(range_match.group(2))
                for n in range(start, end + 1):
                    case_nums.add(n)
            else:
                # 默认范围：咨询实务案例题 Q201-210
                for n in range(201, 211):
                    case_nums.add(n)
                start = 201

            # 提取案例背景文本（括号后面的内容）
            bg_match = re.match(r'^\d+\.\s+案例分析题[（(][^）)]*[）)](.*)', text)
            if bg_match:
                bg_text = bg_match.group(1).strip()
                case_backgrounds.append({
                    "start_num": start,
                    "background_text": bg_text,
                    "range_text": text,
                })
    return case_nums, case_backgrounds


def _clean_case_background(text):
    """清洗案例背景文本，去除题号前缀和格式标记"""
    # 去掉开头的 "201. 案例分析题（...）" 前缀
    text = re.sub(r'^\d+\.\s+案例分析题[（(][^）)]*[）)]\s*', '', text)
    text = text.strip()
    # 压缩多余空白
    text = re.sub(r'\s+', ' ', text)
    return text


def infer_category(source_file):
    """从文件名推断知识板块"""
    if not source_file:
        return "其他"
    for key, cat in CATEGORY_MAP.items():
        if key in source_file:
            return cat
    return "其他"


def parse_docx(filepath):
    """
    解析单个 docx 文件，提取所有题目

    返回: list[dict]
        {
            "id": "md5前12位",
            "source_file": "文件名",
            "index": 题号,
            "type": "single" | "multi" | "judge",
            "question": "题目内容（不含题号）",
            "options": {"A": "选项", "B": "选项", ...},
            "answer": "A" | "AB" | "正确/错误",
            "explanation": "答案解析",
            "md5": "完整md5",
            "category": "基础理论|咨询实务|其他"
        }
    """
    filepath = Path(filepath)
    doc = docx.Document(str(filepath))
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    # 预扫描：检测案例题范围
    case_nums, case_backgrounds = _detect_case_ranges(paragraphs)

    # 题目正则：数字开头，含题型标记
    q_pattern = re.compile(r'^(\d+)\.\s+(.+?)\s*\((单选题|多选题|判断题|不定项选择题)')
    opt_pattern = re.compile(r'^([A-Z])[：:]\s*(.*)')
    ans_pattern = re.compile(r'^正确答案[：:]\s*(.+)')
    exp_pattern = re.compile(r'^答案解析[：:]\s*(.+)')

    questions = []
    current = None
    source_name = filepath.name

    for text in paragraphs:
        if not text:
            # 空行：如果正在收集题目，可能表示题目结束
            continue

        # 检测新题目
        m = q_pattern.match(text)
        if m:
            # 保存上一题
            if current and current.get("question"):
                _finalize_question(current)
                questions.append(current)

            idx = int(m.group(1))
            q_text = m.group(2).strip()
            q_type = m.group(3)

            type_map = {"单选题": "single", "多选题": "multi", "判断题": "judge", "不定项选择题": "indefinite", "案例题": "案例题"}

            # 判断题的题目中可能包含 "正确/错误" 标记，需要清理
            if q_type == "判断题":
                q_text = re.sub(r'[（(]\s*[）)]\s*$', '', q_text).strip()

            current = {
                "source_file": source_name,
                "index": idx,
                "type": type_map[q_type],
                "question": q_text,
                "options": {},
                "answer": "",
                "explanation": "",
                "category": infer_category(source_name),
            }

            # 案例题范围：将题型改为"案例题"
            if idx in case_nums:
                current["type"] = "案例题"

            continue

        if current is None:
            continue

        # 检测选项行
        m = opt_pattern.match(text)
        if m:
            opt_key = m.group(1)
            opt_val = m.group(2).strip()
            current["options"][opt_key] = opt_val
            continue

        # 检测正确答案
        m = ans_pattern.match(text)
        if m:
            current["answer"] = m.group(1).strip()
            continue

        # 检测答案解析
        m = exp_pattern.match(text)
        if m:
            current["explanation"] = m.group(1).strip()
            continue

    # 保存最后一题
    if current and current.get("question"):
        _finalize_question(current)
        questions.append(current)

    return questions


def _finalize_question(q):
    """完善题目数据：生成ID和MD5"""
    # 如判断题无选项，自动生成
    if q["type"] == "judge" and not q["options"]:
        q["options"] = {"A": "正确", "B": "错误"}

    # 判断题答案标准化
    if q["type"] == "judge":
        ans = q["answer"].strip()
        if ans in ("A", "正确", "对", "√", "true", "True"):
            q["answer"] = "A"
        elif ans in ("B", "错误", "错", "×", "false", "False"):
            q["answer"] = "B"

    # 多选题/不定项选择题/案例题答案标准化：去掉空格，字母排序
    if q["type"] in ("multi", "案例题", "indefinite"):
        ans = q["answer"].upper().replace(" ", "").replace("，", "").replace(",", "")
        ans = "".join(sorted(set(ans)))
        q["answer"] = ans

    # 单选题答案标准化
    if q["type"] == "single":
        q["answer"] = q["answer"].strip().upper()

    # 移除答案中的不可见字符（Word 文档导入时常见的零宽空格等格式污染）
    q["answer"] = _sanitize_answer(q["answer"])

    # 生成MD5用于去重
    content = q["question"] + str(sorted(q["options"].items())) + q["answer"]
    md5_hash = hashlib.md5(content.encode("utf-8")).hexdigest()
    q["md5"] = md5_hash
    q["id"] = md5_hash[:12]


def batch_parse(docx_dir, progress_callback=None):
    """
    批量解析目录下所有 docx 文件（递归扫描子目录）

    返回: (list[dict] 所有题目, list[dict] 文件统计, list[dict] 案例背景)
    案例背景格式: {"case_id": str(UUID), "title": str, "start_num": int, "source_file": str}
    """
    docx_dir = Path(docx_dir)
    all_questions = []
    file_stats = []
    all_case_backgrounds = []

    # 递归扫描，过滤临时文件
    files = sorted(f for f in docx_dir.rglob("*.docx") if not f.name.startswith("~$"))
    for i, fpath in enumerate(files):
        try:
            qs = parse_docx(fpath)
            all_questions.extend(qs)
            stats = {
                "file": fpath.name,
                "total": len(qs),
                "single": sum(1 for q in qs if q["type"] == "single"),
                "multi": sum(1 for q in qs if q["type"] == "multi"),
                "judge": sum(1 for q in qs if q["type"] == "judge"),
            }
            file_stats.append(stats)

            # 提取案例背景（从docx直接读取）
            case_bgs = _extract_case_backgrounds_from_docx(fpath)
            all_case_backgrounds.extend(case_bgs)

            if progress_callback:
                progress_callback(i + 1, len(files), fpath.name, stats)
        except Exception as e:
            file_stats.append({"file": fpath.name, "error": str(e)})

    return all_questions, file_stats, all_case_backgrounds


def _extract_case_backgrounds_from_docx(filepath):
    """
    从 docx 文件中提取案例背景信息，返回 case_study 格式的 dict 列表。
    每个 dict: {"case_id": str, "title": str, "start_num": int, "source_file": str, "end_num": int}
    """
    filepath = Path(filepath)
    doc = docx.Document(str(filepath))
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    case_nums, case_bgs_raw = _detect_case_ranges(paragraphs)
    results = []
    for bg in case_bgs_raw:
        # 生成确定性UUID（基于源文件名+起始题号）
        seed = f"{filepath.name}:{bg['start_num']}"
        case_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, seed))
        title = _clean_case_background(bg["range_text"])
        # 确定结束题号
        end_num = max(n for n in case_nums if n >= bg["start_num"]) if case_nums else bg["start_num"] + 9
        results.append({
            "case_id": case_id,
            "title": title,
            "start_num": bg["start_num"],
            "end_num": end_num,
            "source_file": filepath.name,
        })
    return results


if __name__ == "__main__":
    # 测试解析
    import json
    test_file = Path(__file__).parent.parent / "exmbase" / "【基础理论1】心理学导论.docx"
    if test_file.exists():
        qs = parse_docx(test_file)
        print(f"解析完成：{len(qs)} 题")
        for q in qs[:3]:
            print(json.dumps(q, ensure_ascii=False, indent=2))
            print("---")
    else:
        print(f"文件不存在: {test_file}")
